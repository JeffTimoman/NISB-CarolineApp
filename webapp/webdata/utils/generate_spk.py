from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- CREATE DOCUMENT & BASE STYLE ---
document = Document()

# Set page margins (optional, adjust as needed)
for section in document.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

# Base font style
style = document.styles["Normal"]
font = style.font
font.name = "Times New Roman"
font.size = Pt(12)

# Helper to add a blank line
def add_blank_lines(doc, n=1):
    for _ in range(n):
        doc.add_paragraph("")

# --- HEADER (No & Date) ---
p_no = document.add_paragraph()
p_no.add_run("No:  /IIO/year")

p_date = document.add_paragraph()
p_date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p_date.add_run("Jakarta, {MMM YYYY}")

add_blank_lines(document, 1)

# --- ADDRESS BLOCK ---
document.add_paragraph("Kepada Yth,")
document.add_paragraph("PT HARAP ISI")
document.add_paragraph("UP. NAMA UP")
document.add_paragraph("Di Tempat")

add_blank_lines(document, 1)

# --- SUBJECT (Perihal) ---
p_perihal = document.add_paragraph()
r_label = p_perihal.add_run("Perihal: ")
r_label.bold = True
p_perihal.add_run("{Remain Empty}")

add_blank_lines(document, 1)

# --- BODY TEXT ---
body1 = (
    "Sehubungan dengan kebutuhan link komunikasi cabang PT. Bank Central Asia Tbk. (BCA), "
    "kami mohon untuk dilakukan Instalasi Jaringan Komunikasi dengan rincian sebagai berikut :"
)
document.add_paragraph(body1)

add_blank_lines(document, 1)

# --- DETAIL SECTION AS TABLE (Label | Value) ---
details = [
    ("Nama Cabang", "{Lokasi if only 1 ip, Terlampir (bold) if multiple}"),
    (
        "Alamat",
        "{FORMAT: JL.Raya Rancaekek KM 24.5,Kawasan Dwipapuri,Desa Mangunarga,"
        "Kec.Cimanggung,Kab.Sumedang, Jawa Barat\n\n"
        "Lat, Long         : -6.9556222 , 107.7979483 Filled if single, "
        "Terlampir (bold) if multiple}"
    ),
    ("Layanan", "GBR {Nama Provider}"),
    ("Bandwidth", "64 Kbps"),
    ("Kuota", "1 Gb"),
    ("PIC Lokasi", "{Ibu Rara (0857-7536-9737)}"),
]

table = document.add_table(rows=0, cols=2)
table.style = "Table Grid"  # You can change style or remove

for label, value in details:
    row_cells = table.add_row().cells
    row_cells[0].text = label
    row_cells[1].text = value

add_blank_lines(document, 1)

# --- BODY TEXT (PKS REFERENCE) ---
body2 = (
    "Pengerjaan Instalasi Layanan Link GBR di lokasi terlampir hendaknya mengacu pada "
    "Perjanjian Kerja Sama (PKS) antara HARAP ISI INI dan BCA"
)
document.add_paragraph(body2)

add_blank_lines(document, 1)

body3 = "Demikian kami sampaikan dan terima kasih atas kerja sama yang baik."
document.add_paragraph(body3)

add_blank_lines(document, 3)

# --- SIGNATURE BLOCK ---
p_hormat = document.add_paragraph("Hormat Kami,")
add_blank_lines(document, 4)

p_name = document.add_paragraph()
run_name = p_name.add_run("ROBERTUS ADITYA SETIAWAN")
run_name.bold = True

document.add_paragraph("Vice President")

add_blank_lines(document, 2)

document.add_paragraph("By : OLN")

add_blank_lines(document, 2)

# --- LAMPIRAN TITLE ---
p_lampiran = document.add_paragraph()
r_lampiran = p_lampiran.add_run("LAMPIRAN")
r_lampiran.bold = True
r_lampiran.underline = True

add_blank_lines(document, 1)

# --- LAMPIRAN CONTENT (MULTIPLE DATA FORMAT) ---
# Replace the example paragraph blocks with a table of lampiran entries.
# The `lampiran_items` list should contain dicts with keys:
#  - lokasi, alamat, latlong, pic
# This keeps the layout consistent when there are many entries.

# Lampiran items provided by the user
lampiran_items = [
    {
        "lokasi": "SNS JAKARTA",
        "alamat": "JL SNS JAKARTA",
        "latlong": "-1, 1",
        "pic": "Ibu Rara (0812 7473 7779)",
    },
    {
        "lokasi": "SNS JAKARTA 2",
        "alamat": "JL SNS JAKARTA2",
        "latlong": "-0.5, 0.5",
        "pic": "Ibu Rari (0813 7757 9888)",
    },
    {
        "lokasi": "SNS JAKARTA 3",
        "alamat": "JL SNS JAKARTA3",
        "latlong": "-0.25, 0.25",
        "pic": "Ibu Rare (0813 7311 8686)",
    },
]

# Build a single Table Grid with one row per lampiran item
rows = 1 + len(lampiran_items)
cols = 5  # No | Lokasi | Alamat | Lat, Long | PIC Lokasi
lamp_table = document.add_table(rows=rows, cols=cols)
lamp_table.style = "Table Grid"

# Header
hdr = lamp_table.rows[0].cells
hdr[0].text = "No"
hdr[1].text = "Lokasi"
hdr[2].text = "Alamat"
hdr[3].text = "Lat, Long"
hdr[4].text = "PIC Lokasi"

# Populate rows
for idx, item in enumerate(lampiran_items, start=1):
    row = lamp_table.rows[idx].cells
    row[0].text = str(idx)
    row[1].text = item.get("lokasi", "")
    row[2].text = item.get("alamat", "")
    row[3].text = item.get("latlong", "")
    row[4].text = item.get("pic", "")

add_blank_lines(document, 2)

# --- FOOTER TEXT (BCA INFO) ---
footer_text = (
    "PT BANK CENTRAL ASIA TBK\n"
    "Head Office: Menara BCA Grand Indonesia, JI. M. H. Thamrin No. I Jakarta 10310 "
    "Tel. (021) 2358-8000 Fax. (021) 2358-8300"
)
for line in footer_text.split("\n"):
    document.add_paragraph(line)

# --- SAVE DOCUMENT ---
document.save("SPK_Instalasi_Layanan_GBR_Indosat_Template.docx")
