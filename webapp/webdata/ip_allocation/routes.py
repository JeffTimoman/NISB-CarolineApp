from flask import Blueprint, request, url_for, redirect, render_template, flash
from flask import url_for, redirect, request, jsonify, current_app
from werkzeug.utils import secure_filename
from webdata import db, Config
from webdata.models import IpAllocation
from flask import send_file
import io
import os
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import ipaddress

ip_allocation = Blueprint('ip_allocation', __name__)

@ip_allocation.route("/upload_file", methods=["POST"])
def upload_file():
	"""Handle uploaded Excel for IP allocation.

	Expects a file in the form field named 'file'.
	Reads:
	  - start_ip from cell F3
	  - provider from cell F4
	  - details from rows starting at row 7 with columns:
		  C: location, D: address, E: lat, F: long, G: pic_bca, H: phone, I: online_date

	Creates an IpAllocation and associated IpAllocationDetail records.
	Returns JSON with success or error message.
	"""
	from flask import current_app
	import os
	from webdata.models import IpAllocation, IpAllocationDetail
	from webdata.utils.excel import parse_ip_allocation_excel
	# validate file
	if 'file' not in request.files:
		return jsonify({'success': False, 'message': 'No file part in request'}), 400

	file = request.files['file']
	if file.filename == '':
		return jsonify({'success': False, 'message': 'No selected file'}), 400

	orig_filename = file.filename
	# create a safe original filename for extension extraction
	safe_orig = secure_filename(orig_filename)
	# generate a uuid filename to store on disk, keep original extension if present
	import uuid as _uuid
	_, ext = os.path.splitext(safe_orig)
	stored_filename = f"{str(_uuid.uuid4())}{ext}"

	# ensure excel folder exists (use project root from Config)
	excel_folder = os.path.join(os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..')), Config.EXCEL_FOLDER_PATH)
	os.makedirs(excel_folder, exist_ok=True)

	saved_path = os.path.join(excel_folder, stored_filename)
	file.save(saved_path)

	try:
		parsed = parse_ip_allocation_excel(saved_path)
	except Exception as e:
		return jsonify({'success': False, 'message': f'Failed to parse excel: {str(e)}'}), 400

	# Build models
	ipalloc = IpAllocation(
		start_ip=parsed.get('start_ip'),
		provider=parsed.get('provider'),
		original_excel_name=orig_filename,
		excel_filename=stored_filename,
	)

	details = []
	for idx, d in enumerate(parsed.get('details', []), start=1):
		detail = IpAllocationDetail(
			order=idx,
			location=d.get('location'),
			address=d.get('address'),
			lat=d.get('lat'),
			long=d.get('long'),
			pic_bca=d.get('pic_bca'),
			pic_bca_phone=d.get('pic_bca_phone'),
			online_date=d.get('online_date'),
		)
		details.append(detail)

	# attach and commit
	ipalloc.details = details
	db.session.add(ipalloc)
	db.session.commit()

	# Prepare response including parsed extraction so frontend can show details
	response_parsed = {
		'success': True,
		'id': ipalloc.id,
		'message': 'Uploaded and saved successfully',
		'start_ip': parsed.get('start_ip'),
		'provider': parsed.get('provider'),
		'details': [],
	}

	from datetime import datetime
	for d in parsed.get('details', []):
		online = d.get('online_date')
		if isinstance(online, datetime):
			online_val = online.isoformat()
		else:
			online_val = online
		response_parsed['details'].append({
			'location': d.get('location'),
			'address': d.get('address'),
			'lat': d.get('lat'),
			'long': d.get('long'),
			'pic_bca': d.get('pic_bca'),
			'pic_bca_phone': d.get('pic_bca_phone'),
			'online_date': online_val,
		})

	return jsonify(response_parsed)


@ip_allocation.route('/detail', methods=['GET'])
def detail():
	"""View-only detail page for a saved IpAllocation. Query param: id"""
	from flask import render_template, request, abort, url_for
	from webdata.models import IpAllocation
	import ipaddress

	alloc_id = request.args.get('id')
	if not alloc_id:
		abort(400, description='Missing id')

	ipalloc = IpAllocation.query.filter_by(id=alloc_id).first()
	if not ipalloc:
		abort(404, description='IpAllocation not found')

	details = sorted(ipalloc.details, key=lambda d: (d.order or 0))

	# Compute sequential start/end IPs for each detail based on ipalloc.start_ip and each detail.mask
	allocations = []
	try:
		current = ipaddress.IPv4Address(ipalloc.start_ip) if ipalloc.start_ip else None
	except Exception:
		current = None

	for d in details:
		assigned_start = None
		assigned_end = None
		gateway = None
		atm_start = None
		atm_end = None
		subnetmask = None
		prefix_int = None
		if current is not None:
			# determine prefix from mask (expect format like '/29' or '29')
			mask = (d.mask or '').strip()
			if mask.startswith('/'):
				prefix = mask[1:]
			else:
				prefix = mask
			try:
				prefix_int = int(prefix)
				# number of addresses in the block
				block_size = 2 ** (32 - prefix_int) if 0 <= prefix_int <= 32 else None
			except Exception:
				block_size = None

			if block_size:
				assigned_start = ipaddress.IPv4Address(int(current))
				assigned_end = ipaddress.IPv4Address(int(current) + block_size - 1)

				# compute network and hosts using ipaddress to get gateway and usable host ranges
				try:
					network = ipaddress.ip_network(f"{assigned_start}/{prefix_int}", strict=False)
					hosts = list(network.hosts())  # usable hosts
					if hosts:
						gateway = hosts[0]  # first usable as gateway
					if len(hosts) >= 2:
						atm_start = hosts[1]
						atm_end = hosts[-1]
					elif len(hosts) == 1:
						atm_start = hosts[0]
						atm_end = hosts[0]
					subnetmask = str(network.netmask)
				except Exception:
					gateway = None
					atm_start = None
					atm_end = None
					subnetmask = None

				# advance current to next address after this block
				current = ipaddress.IPv4Address(int(current) + block_size)

		allocations.append({
			'detail': d,
			'start_ip': str(assigned_start) if assigned_start is not None else None,
			'end_ip': str(assigned_end) if assigned_end is not None else None,
			'gateway': str(gateway) if gateway is not None else None,
			'atm_start': str(atm_start) if atm_start is not None else None,
			'atm_end': str(atm_end) if atm_end is not None else None,
			'subnetmask': subnetmask,
			'prefix': prefix_int,
		})

	return render_template('ip_allocation/detail.html', ipalloc=ipalloc, details=details, allocations=allocations)


@ip_allocation.route('/edit', methods=['GET', 'POST'])
def edit():
	"""Edit page for IpAllocation. GET shows editable form; POST saves changes."""
	from flask import render_template, request, abort, redirect, url_for, flash
	from webdata.models import IpAllocation
	from datetime import datetime

	alloc_id = request.args.get('id') or request.form.get('ipalloc_id')
	if not alloc_id:
		abort(400, description='Missing id')

	ipalloc = IpAllocation.query.filter_by(id=alloc_id).first()
	if not ipalloc:
		abort(404, description='IpAllocation not found')

	if request.method == 'POST':
		# update fields for each detail
		for d in ipalloc.details:
			prefix = f"{d.id}"
			d.location = request.form.get(f'location-{prefix}')
			d.address = request.form.get(f'address-{prefix}')
			d.lat = request.form.get(f'lat-{prefix}')
			d.long = request.form.get(f'long-{prefix}')
			d.pic_bca = request.form.get(f'pic_bca-{prefix}')
			d.pic_bca_phone = request.form.get(f'pic_bca_phone-{prefix}')
			mask = request.form.get(f'mask-{prefix}')
			if mask is not None:
				d.mask = mask
			d.bandwidth = request.form.get(f'bandwidth-{prefix}')
			d.quota = request.form.get(f'quota-{prefix}')
			# order can be provided by the form (order-<id>) when rows are reordered in the UI
			order_raw = request.form.get(f'order-{prefix}')
			if order_raw:
				try:
					d.order = int(order_raw)
				except Exception:
					pass

			online_raw = request.form.get(f'online_date-{prefix}')
			if online_raw:
				try:
					dt = datetime.fromisoformat(online_raw)
				except Exception:
					try:
						dt = datetime.strptime(online_raw, '%Y-%m-%dT%H:%M')
					except Exception:
						dt = None
				d.online_date = dt
			else:
				d.online_date = None

		db.session.commit()
		flash('Changes saved', 'success')
		return redirect(url_for('ip_allocation.detail') + '?id=' + alloc_id)

	# GET
	details = sorted(ipalloc.details, key=lambda d: (d.order or 0))
	return render_template('ip_allocation/edit.html', ipalloc=ipalloc, details=details)


@ip_allocation.route('/api/ip_allocations', methods=['GET'])
def api_ip_allocations():
	"""Return paginated IpAllocation items for the history UI.

	Query params: page (int), per_page (int)
	"""
	page = request.args.get('page', 1, type=int)
	per_page = request.args.get('per_page', 5, type=int)
	pagination = IpAllocation.query.order_by(IpAllocation.created_at.desc()).paginate(page=page, per_page=per_page, error_out=False)
	items = [{
		'id': a.id,
		'name': a.name,
		'start_ip': a.start_ip,
		'provider': a.provider,
		'original_excel_name': getattr(a, 'original_excel_name', None),
		'excel_filename': getattr(a, 'excel_filename', None),
		'created_at': a.created_at.strftime('%Y-%m-%d %H:%M:%S') if getattr(a, 'created_at', None) else None,
		'updated_at': a.updated_at.strftime('%Y-%m-%d %H:%M:%S') if getattr(a, 'updated_at', None) else None,
	} for a in pagination.items]

	return jsonify({
		'items': items,
		'page': pagination.page,
		'pages': pagination.pages,
		'has_next': pagination.has_next,
		'has_prev': pagination.has_prev,
		'next_num': pagination.next_num,
		'prev_num': pagination.prev_num,
		'total': pagination.total
	})

@ip_allocation.route('/api/ip_allocations/<id>/name', methods=['PUT'])
def update_ip_allocation_name(id):
	data = request.get_json() or {}
	new_name = data.get('name')
	if not new_name or not str(new_name).strip():
		return jsonify({'error': 'Name is required'}), 400

	alloc = IpAllocation.query.get(id)
	if not alloc:
		return jsonify({'error': 'Not found'}), 404

	alloc.name = str(new_name).strip()
	db.session.commit()
	return jsonify({'success': True, 'name': alloc.name})





































@ip_allocation.route('/generate_spk', methods=['GET'])
def generate_spk():
    """Generate SPK .docx for a given IpAllocation id and return as attachment."""
    alloc_id = request.args.get('id')
    if not alloc_id:
        return jsonify({'error': 'Missing id'}), 400

    ipalloc = IpAllocation.query.filter_by(id=alloc_id).first()
    if not ipalloc:
        return jsonify({'error': 'Not found'}), 404

    # gather details ordered by "order"
    details = sorted(ipalloc.details, key=lambda d: (d.order or 0))

    # create document
    doc = Document()

    # === PAGE SETUP: A4 + margins + header/footer distance ===
    for section in doc.sections:
        # A4 size
        section.page_width = Cm(21)     # 210 mm
        section.page_height = Cm(29.7)  # 297 mm

        # Margins (inches -> cm)
        section.top_margin = Cm(1.4224)    # 0.56"
        section.left_margin = Cm(1.524)    # 0.6"
        section.bottom_margin = Cm(0.4826) # 0.19"
        section.right_margin = Cm(2.1082)  # 0.83"

        # Header / footer distance from edge (0.5")
        section.header_distance = Cm(1.27)
        section.footer_distance = Cm(1.27)

    # Default body style: Verdana 11
    style = doc.styles['Normal']
    base_font = style.font
    base_font.name = 'Verdana'
    base_font.size = Pt(11)

    # helpers
    def set_run_font(run, name, size_pt=None, bold=None, underline=None):
        if name:
            run.font.name = name
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if bold is not None:
            run.font.bold = bold
        if underline is not None:
            run.font.underline = underline

    def set_paragraph_font(paragraph, name, size_pt):
        for r in paragraph.runs:
            set_run_font(r, name, size_pt)

    def make_tight(paragraph):
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    # === HEADER: logo + No/Date line ===
    now = datetime.now()
    year = now.year
    month_str = now.strftime('%b %Y')  # "Nov 2025"

    section = doc.sections[0]
    usable_width = section.page_width - section.left_margin - section.right_margin

    # Header logo (on all pages)
    header = section.header
    if header.paragraphs:
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = ""
    else:
        header_paragraph = header.add_paragraph()
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_run = header_paragraph.add_run()
    # adjust width as needed
    # Resolve static image path reliably using the Flask app root
    image_path = os.path.join(current_app.root_path, 'static', 'img', 'bca.png')
    if os.path.exists(image_path):
        try:
            logo_run.add_picture(image_path, width=Inches(1.6))
        except Exception as e:
            current_app.logger.warning(f"Failed to add logo picture to document: {e}")
    else:
        current_app.logger.warning(f"Logo image not found at {image_path}; skipping logo in document.")
    # add one blank line after the header logo
    try:
        header_paragraph.add_run().add_break()
    except Exception:
        # fallback: add an empty paragraph in the header
        header.add_paragraph('')

    # No / Date line in body (not in header)
    p_header = doc.add_paragraph()
    pf = p_header.paragraph_format
    pf.tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

    r_left = p_header.add_run(f"No:     /IIO/{year}")
    set_run_font(r_left, 'Verdana', 11)

    p_header.add_run("\t")

    r_right = p_header.add_run(f"Jakarta, {month_str}")
    set_run_font(r_right, 'Verdana', 11)

    # small blank after header text
    doc.add_paragraph('')

    # === ADDRESS BLOCK (Arial 11, NO empty lines between them) ===
    p = doc.add_paragraph("Kepada Yth,")
    set_paragraph_font(p, 'Arial', 11)
    make_tight(p)

    p = doc.add_paragraph("PT HARAP ISI")
    set_paragraph_font(p, 'Arial', 11)
    make_tight(p)

    p = doc.add_paragraph("UP. NAMA UP")
    set_paragraph_font(p, 'Arial', 11)
    make_tight(p)

    p = doc.add_paragraph("Di Tempat")
    set_paragraph_font(p, 'Arial', 11)
    make_tight(p)

    # one blank line after Di Tempat
    doc.add_paragraph('')

    # === PERIHAL directly after address block ===
    p_perihal = doc.add_paragraph()
    r_label = p_perihal.add_run('Perihal: ')
    set_run_font(r_label, 'Verdana', 11, bold=True, underline=True)
    # value remains empty

    # NO blank line between Perihal and Sehubungan (keep default spacing)

    # === BODY INTRO (Sehubungan...) – Arial 11 ===
    p_body1 = doc.add_paragraph()
    make_tight(p_body1)

    r1 = p_body1.add_run(
        "Sehubungan dengan kebutuhan link komunikasi cabang PT. Bank Central Asia Tbk. (BCA), "
        "kami mohon untuk dilakukan "
    )
    set_run_font(r1, 'Arial', 11)

    r2 = p_body1.add_run("Instalasi Jaringan Komunikasi")
    set_run_font(r2, 'Arial', 11, bold=True)

    r3 = p_body1.add_run(" dengan rincian sebagai berikut :")
    set_run_font(r3, 'Arial', 11)

    # one blank line BEFORE table
    doc.add_paragraph('')

    # === DETAIL TABLE (main table) ===
    if len(details) == 1:
        d0 = details[0]
        nama_cabang = d0.location or ''
        alamat = d0.address or ''
        latlong_line = ""
        if d0.lat or d0.long:
            lat = d0.lat or ''
            lon = d0.long or ''
            latlong_line = f"Lat, Long         : {lat} , {lon}"

        bandwidth_value = d0.bandwidth or "64 Kbps"
        quota_value = d0.quota or "1 Gb"

        alamat_cell_value = alamat
        if latlong_line:
            alamat_cell_value = alamat_cell_value + '\n' + latlong_line

        nama_cabang_value = nama_cabang
        nama_cabang_bold = False
        alamat_value_bold = False
    else:
        nama_cabang_value = 'Terlampir'
        nama_cabang_bold = True
        alamat_cell_value = 'Terlampir'
        alamat_value_bold = True
        bandwidth_value = "64 Kbps"
        quota_value = "1 Gb"

    pic_lokasi_value = "Terlampir"

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    # indent table 0.5 cm from left (table property tblInd uses dxa units)
    try:
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        tblInd = tblPr.find(qn('w:tblInd'))
        if tblInd is None:
            tblInd = OxmlElement('w:tblInd')
            tblPr.append(tblInd)
        # 1 cm ~= 567 dxa, so 0.8 cm ~= 454
        tblInd.set(qn('w:w'), str(int(0.8 * 567)))
        tblInd.set(qn('w:type'), 'dxa')
    except Exception:
        current_app.logger.warning('Failed to set table left indent')

    # center the table within the page width
    try:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
    except Exception:
        pass

    # Main table columns: 1.2" + 5.64"
    left_col_width = Inches(1.2)
    right_col_width = Inches(5.64)

    def add_row(label, value, bold_value=False):
        row = table.add_row()
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Inches(0.2)  # minimal height (user requested 0.2")

        # Left cell (label)
        c0 = row.cells[0]
        c0.width = left_col_width
        c0.text = ""
        p0 = c0.paragraphs[0]
        # left-align label text but vertically center the cell
        try:
            p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except Exception:
            pass
        try:
            c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            pass
        r0 = p0.add_run(label)
        set_run_font(r0, 'Arial', 11, bold=False)

        # Right cell (value)
        c1 = row.cells[1]
        c1.width = right_col_width
        c1.text = ""
        if value is None:
            value = ""
        p1 = c1.paragraphs[0]
        try:
            c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        except Exception:
            pass
        lines = str(value).split('\n')
        for i, line in enumerate(lines):
            r_line = p1.add_run(line)
            set_run_font(r_line, 'Arial', 11, bold=bold_value)
            if i != len(lines) - 1:
                nl = p1.add_run('\n')
                set_run_font(nl, 'Arial', 11, bold=bold_value)

    add_row('Nama Cabang', nama_cabang_value, bold_value=nama_cabang_bold)
    add_row('Alamat', alamat_cell_value, bold_value=alamat_value_bold)
    add_row('Layanan', f"GBR {ipalloc.provider or ''}")
    add_row('Bandwidth', bandwidth_value)
    add_row('Kuota', quota_value)
    add_row('PIC Lokasi', pic_lokasi_value, bold_value=True)

    # one blank line AFTER table
    doc.add_paragraph('')

    # === PARAGRAPHS UNDER THE TABLE ===

    # Pengerjaan...
    p_body2 = doc.add_paragraph()
    make_tight(p_body2)
    r4 = p_body2.add_run(
        "Pengerjaan Instalasi Layanan Link GBR di lokasi terlampir hendaknya mengacu pada "
    )
    set_run_font(r4, 'Arial', 11)

    r5 = p_body2.add_run("Perjanjian Kerja Sama (PKS)")
    set_run_font(r5, 'Arial', 11, bold=True)

    r6 = p_body2.add_run(" antara HARAP ISI INI dan BCA")
    set_run_font(r6, 'Arial', 11)

    # Demikian...
    p_close = doc.add_paragraph(
        "Demikian kami sampaikan dan terima kasih atas kerja sama yang baik."
    )
    set_paragraph_font(p_close, 'Arial', 11)
    make_tight(p_close)

    # One blank line before signature
    doc.add_paragraph('')

    # === SIGNATURE BLOCK ===
    p_sig1 = doc.add_paragraph("Hormat Kami,")
    set_paragraph_font(p_sig1, 'Arial', 11)

    # signature space
    doc.add_paragraph('')
    doc.add_paragraph('')

    # Name line
    p_name = doc.add_paragraph()
    r_name = p_name.add_run("ROBERTUS ADITYA SETIAWAN")
    set_run_font(r_name, 'Arial', 11, bold=True, underline=True)

    # Title line
    p_title = doc.add_paragraph("Vice President")
    for r in p_title.runs:
        set_run_font(r, 'Arial', 11, bold=True)

    doc.add_paragraph('')

    # BY: OLN Verdana 7
    p_by = doc.add_paragraph("By : OLN")
    for r in p_by.runs:
        set_run_font(r, 'Verdana', 7)

    # small gap before Lampiran page
    doc.add_paragraph('')

    # === PAGE BREAK BEFORE LAMPIRAN PAGE ===
    doc.add_page_break()

    # === LAMPIRAN HEADING ON NEW PAGE ===
    p_lamp = doc.add_paragraph()
    p_lamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lamp = p_lamp.add_run('LAMPIRAN')
    set_run_font(r_lamp, 'Verdana', 11, bold=True)
    doc.add_paragraph('')

    # === LAMPIRAN TABLES (details) ===
    # Font: Arial 10, columns: 1.06" / 5.32"
    if len(details) > 1:
        for d in details:
            detail_table = doc.add_table(rows=0, cols=2)
            detail_table.style = 'Table Grid'
            detail_table.autofit = False
            detail_table.allow_autofit = False

            lamp_left_width = Inches(1.06)
            lamp_right_width = Inches(5.32)

            def add_label_row(label, value):
                row = detail_table.add_row()
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
                row.height = Inches(0.11)

                c0 = row.cells[0]
                c0.width = lamp_left_width
                c0.text = ""
                p0 = c0.paragraphs[0]
                r0 = p0.add_run(label)
                set_run_font(r0, 'Arial', 10)

                c1 = row.cells[1]
                c1.width = lamp_right_width
                c1.text = ""
                p1 = c1.paragraphs[0]
                lines = (value or "").split('\n')
                for i, line in enumerate(lines):
                    r_line = p1.add_run(line)
                    set_run_font(r_line, 'Arial', 10)
                    if i != len(lines) - 1:
                        nl = p1.add_run('\n')
                        set_run_font(nl, 'Arial', 10)

            # Lokasi
            add_label_row('Lokasi', d.location or "")

            # Alamat + Lat/Long combined in one cell
            addr_lines = []
            if d.address:
                addr_lines.append(d.address)

            if d.lat or d.long:
                lat = d.lat or ''
                lon = d.long or ''
                addr_lines.append(f"Lat, Long         : {lat} , {lon}")

            alamat_combined = "\n".join(addr_lines)
            add_label_row('Alamat', alamat_combined)

            # PIC Lokasi
            pic = f"{d.pic_bca or ''} {('(' + d.pic_bca_phone + ')') if d.pic_bca_phone else ''}".strip()
            add_label_row('PIC Lokasi', pic)

            doc.add_paragraph('')

    # === REAL FOOTER (all pages) ===
    footer = doc.sections[0].footer

    # First line: PT BANK CENTRAL ASIA TBK
    if footer.paragraphs:
        f1 = footer.paragraphs[0]
        f1.text = ""
    else:
        f1 = footer.add_paragraph()
    r_f1 = f1.add_run("PT BANK CENTRAL ASIA TBK")
    set_run_font(r_f1, 'Times New Roman', 13, bold=True)
    f1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # apply a full-width underline (bottom border) to the paragraph
    def set_paragraph_bottom_border(paragraph, val='single', sz=6, space=1, color='000000'):
        p = paragraph._p
        # ensure pPr exists
        pPr = p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p.insert(0, pPr)

        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is None:
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)

        # remove existing bottom if present
        for child in list(pBdr):
            if child.tag == qn('w:bottom'):
                pBdr.remove(child)

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), val)
        bottom.set(qn('w:sz'), str(sz))
        bottom.set(qn('w:space'), str(space))
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)

    try:
        set_paragraph_bottom_border(f1, val='single', sz=6, space=1, color='000000')
    except Exception:
        # if anything fails, leave footer as-is (no crash)
        current_app.logger.warning('Failed to apply paragraph bottom border to footer line')

    # Second line: Head Office...
    if len(footer.paragraphs) > 1:
        f2 = footer.paragraphs[1]
        f2.text = ""
    else:
        f2 = footer.add_paragraph()
    r_f2 = f2.add_run(
        "Head Office: Menara BCA Grand Indonesia, JI. M. H. Thamrin No. I "
        "Jakarta 10310 Tel. (021) 2358-8000 Fax. (021) 2358-8300"
    )
    set_run_font(r_f2, 'Times New Roman', 8, bold=False, underline=False)
    f2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Save to BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"SPK_Instalasi_{ipalloc.id}.docx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )